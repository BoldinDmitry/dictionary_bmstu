import datetime
import inspect
import io
import pathlib
import sys

from docx import Document

from dictionary_generator import exercises_abs
from dictionary_generator.exercises import *


def fill_row(row, text_fields):
    for i, text in enumerate(text_fields):
        row[i].text = text


def generate(
        name: str,
        date_of_birth: str,
        group: str,
        weight: str,
        height: str,
        start: datetime,
        end: datetime,
        frequency: int,
) -> io.BytesIO:
    """
    Function for generating word file with table
    :param name: student info
    :param date_of_birth: student info
    :param group: student info
    :param weight: student info
    :param height: student info
    :param start: start date of exercises
    :param end: end date of exercises
    :param frequency: how often did student take an exercises
    :return: buffer with docx file
    """
    document = Document()
    document.add_heading("Дневник самоподготовки", 0)

    table = document.add_table(rows=1, cols=2)

    row = table.add_row().cells
    row[0].text = "ФИО"
    row[1].text = name

    row = table.add_row().cells
    row[0].text = "Дата рождения"
    row[1].text = date_of_birth

    row = table.add_row().cells
    row[0].text = "Группа"
    row[1].text = group

    row = table.add_row().cells
    row[0].text = "Вес"
    row[1].text = weight

    row = table.add_row().cells
    row[0].text = "Рост"
    row[1].text = height

    weekdays = [
        "Понедельник",
        "Вторник",
        "Среда",
        "Четверг",
        "Пятница",
        "Суббота",
        "Воскресенье",
    ]
    day_time = ["утро", "день"]

    date_generated = [
        start + datetime.timedelta(days=x)
        for x in range(0, (end - start).days, frequency)
    ]

    all_exercises = []
    for name, obj in inspect.getmembers(sys.modules[__name__]):
        if inspect.isclass(obj) and obj != exercises_abs.ExerciseAbstract:
            all_exercises.append(obj())

    table = document.add_table(rows=1, cols=5)
    table.style = "Table Grid"

    hdr_cells = table.rows[0].cells

    column_names = ["День недели, число, месяц, год, время занятия", "Содержания физкультурного занятия", "Пульс",
                    "Самочувствие", "Желание заниматься"]

    fill_row(hdr_cells, column_names)

    for date in date_generated:
        string_date = ", ".join(
            [
                weekdays[date.weekday()],
                date.strftime("%d.%m.%Y"),
                random.choice(day_time),
            ]
        )

        random.shuffle(all_exercises)

        exercises_count = 4
        description, pulse, state_of_health = "", "", ""

        for i in range(exercises_count):
            description += all_exercises[i].description() + "\n"
            pulse += str(all_exercises[i].pulse()) + "\n"
            state_of_health += all_exercises[i].state_of_health() + "\n"

        row_cells = table.add_row().cells

        column_data = [string_date, description, pulse, state_of_health, "+"]

        fill_row(row_cells, column_data)

    file_stream = io.BytesIO()
    document.save(file_stream)
    file_stream.seek(0)

    return file_stream


def main():
    name = input("ФИО:\n")
    date_of_birth = input("Дата рождения:\n")
    group = input("Группа:\n")
    weight = input("Вес:\n")
    height = input("Рост:\n")

    document_name = input("название документа:\n")

    while True:
        start_str = input("дата начала занятий, формат: %d.%m.%Y:\n")
        try:
            start = datetime.datetime.strptime(start_str, "%d.%m.%Y")
        except ValueError:
            print("некоректная дата")
        else:
            break

    while True:
        end_str = input("дата конца занятий, формат: %d.%m.%Y:\n")
        try:
            end = datetime.datetime.strptime(end_str, "%d.%m.%Y")
        except ValueError:
            print("некоректная дата")
        else:
            break

    while True:
        try:
            frequency = int(input("раз в сколько дней были занятия?\n"))
        except ValueError:
            print("некоректное число")
        else:
            break

    document = generate(
        name=name,
        date_of_birth=date_of_birth,
        group=group,
        weight=weight,
        height=height,
        start=start,
        end=end,
        frequency=frequency,
    )

    here = pathlib.Path(__file__).parent
    document_path = here / (document_name + ".docx")

    with open(document_path, "wb") as w:
        w.write(document.getbuffer().tobytes())

    print("Документ успешно сохранен в ", document_path)


if __name__ == "__main__":
    main()
