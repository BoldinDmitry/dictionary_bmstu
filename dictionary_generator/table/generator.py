import datetime
import random
import io
from docx import Document

from dictionary_generator.exercises.generator import Exercises


class Generator:
    weekdays = [
        "Понедельник",
        "Вторник",
        "Среда",
        "Четверг",
        "Пятница",
        "Суббота",
        "Воскресенье",
    ]

    column_names = [
        "День недели, число, месяц, год, время занятия",
        "Содержания физкультурного занятия",
        "Пульс",
        "Самочувствие",
        "Желание заниматься",
    ]

    day_time = ["утро", "день"]

    def __init__(self):
        self.exercises = Exercises()

    @staticmethod
    def fill_row(row, text_fields):
        for i, text in enumerate(text_fields):
            row[i].text = text

    @staticmethod
    def fill_table(table, data):
        for row_data in data:
            row = table.add_row().cells
            Generator.fill_row(row, row_data)

    def generate_table(self, table, date_generated):
        hdr_cells = table.rows[0].cells

        self.fill_row(hdr_cells, self.column_names)

        exercises_count = 4
        for date in date_generated:
            string_date = ", ".join(
                [
                    self.weekdays[date.weekday()],
                    date.strftime("%d.%m.%Y"),
                    random.choice(self.day_time),
                ]
            )

            row_cells = table.add_row().cells

            description, pulse, state_of_health = self.exercises.get_random_exercises(
                exercises_count
            )

            column_data = [string_date, description, pulse, state_of_health, "+"]

            self.fill_row(row_cells, column_data)

    def generate(
            self,
            name: str,
            date_of_birth: str,
            group: str,
            weight: str,
            height: str,
            start: datetime,
            end: datetime,
            frequency: int,
    ) -> io.BytesIO:
        """
        Function for generating word file with table
        :param name: student info
        :param date_of_birth: student info
        :param group: student info
        :param weight: student info
        :param height: student info
        :param start: start date of exercises
        :param end: end date of exercises
        :param frequency: how often did student take an exercises
        :return: buffer with docx file
        """
        document = Document()
        document.add_heading("Дневник самоподготовки", 0)

        table = document.add_table(rows=1, cols=2)

        rows_data = [
            ["ФИО", name],
            ["Дата рождения", date_of_birth],
            ["Группа", group],
            ["Вес", weight],
            ["Рост", height],
        ]

        Generator.fill_table(table, rows_data)

        date_generated = [
            start + datetime.timedelta(days=x)
            for x in range(0, (end - start).days, frequency)
        ]

        table = document.add_table(rows=1, cols=5)
        table.style = "Table Grid"

        self.generate_table(table, date_generated)

        file_stream = io.BytesIO()
        document.save(file_stream)
        file_stream.seek(0)

        return file_stream
