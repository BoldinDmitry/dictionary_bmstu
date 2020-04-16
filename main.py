import datetime
import inspect
import sys

from docx import Document

import exercises_abs
from exercises import *


def random_day(x):
    one_day = datetime.timedelta(days=1)
    return random.choice([x - one_day, x, x + one_day])


if __name__ == "__main__":
    document = Document()
    document.add_heading("Дневник самоподготовки", 0)

    name = input("ФИО:\n")
    date_of_birth = input("Дата рождения:\n")
    group = input("Группа:\n")
    weight = input("Вес:\n")
    height = input("Рост:\n")

    table = document.add_table(rows=1, cols=2)

    row = table.add_row().cells
    row[0].text = "ФИО"
    row[1].text = name

    row = table.add_row().cells
    row[0].text = "Дата рождения"
    row[1].text = date_of_birth

    row = table.add_row().cells
    row[0].text = "Группа"
    row[1].text = group

    row = table.add_row().cells
    row[0].text = "Вес"
    row[1].text = weight

    row = table.add_row().cells
    row[0].text = "Рост"
    row[1].text = height

    start_str = input("дата начала занятий, формат: %d.%m.%Y:\n")
    end_str = input("дата конца занятий, формат: %d.%m.%Y:\n")

    start = datetime.datetime.strptime(start_str, "%d.%m.%Y")
    end = datetime.datetime.strptime(end_str, "%d.%m.%Y")

    frequency = 4

    weekdays = ["Понедельник", "Вторник", "Среда", "Четверг", "Пятница", "Суббота", "Воскресенье"]
    day_time = ["утро", "день"]

    date_generated = [
        random_day(start + datetime.timedelta(days=x)) for x in range(0, (end - start).days, frequency)]

    all_exercises = []
    for name, obj in inspect.getmembers(sys.modules[__name__]):
        if inspect.isclass(obj) and obj != exercises_abs.ExerciseAbstract:
            all_exercises.append(obj())

    table = document.add_table(rows=1, cols=5)
    table.style = "Table Grid"

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "День недели, число, месяц, год, время занятия"
    hdr_cells[1].text = "Содержания физкультурного занятия"
    hdr_cells[2].text = "Пульс"
    hdr_cells[3].text = "Самочувствие"
    hdr_cells[4].text = "Желание заниматься"

    for date in date_generated:
        string_date = ", ".join([weekdays[date.weekday()], date.strftime("%d.%m.%Y"), random.choice(day_time)])

        random.shuffle(all_exercises)

        exercises_count = 4
        description, pulse, state_of_health = "", "", ""

        for i in range(exercises_count):
            description += all_exercises[i].description() + "\n"
            pulse += str(all_exercises[i].pulse()) + "\n"
            state_of_health += all_exercises[i].state_of_health() + "\n"

        row_cells = table.add_row().cells

        row_cells[0].text = string_date
        row_cells[1].text = description
        row_cells[2].text = pulse
        row_cells[3].text = state_of_health
        row_cells[4].text = "+"

    document_name = input("название документа:\n")
    document.save(document_name + ".docx")
